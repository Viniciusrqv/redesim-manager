"""
gerador_distrato.py
-------------------
Gera Word + PDF de TERMO DE DISTRATO de prestação de serviços
contábeis pré-preenchido a partir dos dados de uma empresa cadastrada
no app (ou do cache de consulta CNPJ).

Suporta 3 iniciativas:
  - 'consensual'  : ambas as partes concordam (caso mais comum)
  - 'cliente'     : iniciativa do cliente (mudou de contábil etc.)
  - 'escritorio'  : iniciativa da CSM (inadimplência, atrito etc.)

Uso:
    from utils.gerador_distrato import gerar_distrato

    pdf_path = gerar_distrato(
        dados_empresa={
            "razao_social": "MUST FOOD ALIMENTOS E BEBIDAS LTDA",
            "cnpj": "37997050000103",
            "endereco": "...",
            "socios": [{"nome": "...", "cpf": "..."}],
        },
        iniciativa="consensual",
        data_efeito="2026-06-15",
        motivo="Cliente migrou para outro escritório contábil.",
        pasta_destino="/sessions/.../mnt/LICENÇAS/distratos/",
    )
"""
from __future__ import annotations

import os
import subprocess
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# Helpers
# ============================================================
def _formatar_cnpj(cnpj: str) -> str:
    c = re.sub(r"\D", "", cnpj or "")
    if len(c) != 14:
        return cnpj or ""
    return f"{c[:2]}.{c[2:5]}.{c[5:8]}/{c[8:12]}-{c[12:]}"


def _formatar_cpf(cpf: str) -> str:
    c = re.sub(r"\D", "", cpf or "")
    if len(c) != 11:
        return cpf or ""
    return f"{c[:3]}.{c[3:6]}.{c[6:9]}-{c[9:]}"


def _formatar_data(d: str) -> str:
    """ISO YYYY-MM-DD → DD/MM/YYYY"""
    if not d:
        return "____/____/______"
    try:
        dt = datetime.strptime(d[:10], "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return d


def _data_extenso(d: str) -> str:
    """ISO YYYY-MM-DD → 'X de mês de AAAA'"""
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    try:
        dt = datetime.strptime(d[:10], "%Y-%m-%d")
        return f"{dt.day} de {meses[dt.month - 1]} de {dt.year}"
    except Exception:
        return "____ de ____________ de 20____"


def _endereco_resumido(empresa: dict) -> str:
    end = empresa.get("endereco") or {}
    if isinstance(end, str):
        # banco antigo: endereco texto livre
        return end
    partes = [
        " ".join([
            (end.get("logradouro") or "").strip(),
            (end.get("numero") or "").strip(),
        ]).strip(),
        (end.get("complemento") or "").strip(),
        (end.get("bairro") or "").strip(),
        f"{(end.get('municipio') or '').strip()}/"
        f"{(end.get('uf') or '').strip()}".strip("/"),
        (end.get("cep") or "").strip(),
    ]
    return ", ".join(p for p in partes if p)


# ============================================================
# Conteudo dinamico por iniciativa
# ============================================================
def _clausulas_iniciativa(iniciativa: str) -> dict:
    """Devolve textos de cláusulas conforme a iniciativa do distrato."""
    if iniciativa == "cliente":
        return {
            "considerando": (
                "Considerando que o CONTRATANTE manifestou interesse na "
                "rescisão do contrato de prestação de serviços contábeis "
                "firmado com a CONTRATADA, por motivos próprios e "
                "alheios à qualidade dos serviços prestados;"
            ),
            "clausula_motivacao": (
                "A presente rescisão se dá por INICIATIVA EXCLUSIVA DO "
                "CONTRATANTE, ficando registrado que a CONTRATADA "
                "cumpriu todas as obrigações contratuais até a data de "
                "encerramento."
            ),
        }
    if iniciativa == "escritorio":
        return {
            "considerando": (
                "Considerando que a CONTRATADA, no exercício regular de "
                "seu direito de não renovação contratual, optou por não "
                "prosseguir com a prestação dos serviços contábeis "
                "anteriormente acordados;"
            ),
            "clausula_motivacao": (
                "A presente rescisão se dá por INICIATIVA DA CONTRATADA, "
                "tendo sido comunicado o CONTRATANTE com a antecedência "
                "necessária para a transição contábil regular."
            ),
        }
    # default = consensual
    return {
        "considerando": (
            "Considerando que ambas as partes, de comum acordo, decidiram "
            "pela rescisão amigável do contrato de prestação de serviços "
            "contábeis anteriormente firmado;"
        ),
        "clausula_motivacao": (
            "A presente rescisão se dá de FORMA CONSENSUAL entre as "
            "partes, sem ônus de qualquer natureza para nenhuma delas, "
            "ficando ambas livres e desobrigadas das cláusulas do "
            "contrato originalmente celebrado, a partir da data de "
            "efeito acima registrada."
        ),
    }


# ============================================================
# Funcao principal
# ============================================================
def gerar_distrato(
    *,
    dados_empresa: dict,
    iniciativa: str = "consensual",
    data_efeito: str | None = None,
    motivo: str | None = None,
    pasta_destino: str,
    gerar_pdf: bool = True,
) -> dict:
    """Gera Word (e opcionalmente PDF) do distrato pre-preenchido.

    Retorna {"docx": path, "pdf": path_or_None, "filename_base": str}.
    """
    if iniciativa not in ("consensual", "cliente", "escritorio"):
        iniciativa = "consensual"

    if not data_efeito:
        data_efeito = datetime.now().date().isoformat()

    # Garante pasta
    Path(pasta_destino).mkdir(parents=True, exist_ok=True)

    # Nome do arquivo: distrato_RAZAO_SOCIAL_AAAAMMDD
    razao = dados_empresa.get("razao_social", "cliente")
    razao_safe = re.sub(r"[^A-Z0-9]+", "_",
                        razao.upper())[:40].strip("_")
    base = f"distrato_{razao_safe}_{data_efeito.replace('-', '')}"
    docx_path = os.path.join(pasta_destino, f"{base}.docx")
    pdf_path = os.path.join(pasta_destino, f"{base}.pdf")

    # ============ Cria o Word ============
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Estilo padrão
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Cabeçalho CSM ---
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("CSM CONTABILIDADE EMPRESARIAL")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xA7, 0x31, 0x27)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Rua Sussumo Yoshimoto, 50 — Jd. Ipê — Cotia/SP — "
        "CEP 06716-150  |  (11) 4616-5887  |  www.csm.com.br"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    # --- Titulo ---
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(
        "TERMO DE DISTRATO DE PRESTAÇÃO DE SERVIÇOS CONTÁBEIS"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xA7, 0x31, 0x27)

    doc.add_paragraph()

    # --- QUALIFICAÇÃO DAS PARTES ---
    h2 = doc.add_paragraph()
    r = h2.add_run("QUALIFICAÇÃO DAS PARTES")
    r.bold = True
    r.font.color.rgb = RGBColor(0xA7, 0x31, 0x27)
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run("CONTRATADA: ").bold = True
    p.add_run(
        "CSM CONTABILIDADE EMPRESARIAL LTDA, inscrita no CNPJ sob "
        "o nº [CNPJ DA CSM AQUI — preencher], com sede na Rua Sussumo "
        "Yoshimoto, 50, Jd. Ipê, Cotia/SP, CEP 06716-150, neste ato "
        "representada por seu(s) sócio(s) administrador(es) na forma "
        "do contrato social."
    )

    doc.add_paragraph()

    # Dados do cliente
    razao_cliente = dados_empresa.get("razao_social", "—")
    cnpj_cliente = _formatar_cnpj(dados_empresa.get("cnpj", ""))
    end_cliente = _endereco_resumido(dados_empresa)

    p = doc.add_paragraph()
    p.add_run("CONTRATANTE: ").bold = True
    p.add_run(
        f"{razao_cliente}, inscrita no CNPJ sob o nº "
        f"{cnpj_cliente or '_______________________'}, "
        f"com sede em {end_cliente or '________________________________'}, "
        f"neste ato representada na forma de seu contrato social."
    )

    # Sócios (se houver)
    socios = dados_empresa.get("socios") or []
    if socios:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Sócios(s) qualificado(s) no contrato social: ").bold = True
        for s in socios[:5]:  # max 5 pra não inflar
            nome = s.get("nome") or "—"
            qual = s.get("qualificacao") or "Sócio(a)"
            cpf = s.get("cpf")
            cpf_fmt = (f", CPF {_formatar_cpf(cpf)}" if cpf else "")
            doc.add_paragraph(
                f"• {nome} ({qual}{cpf_fmt})",
                style="Normal",
            )

    # --- CONSIDERANDOS ---
    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r = h2.add_run("CONSIDERANDOS")
    r.bold = True
    r.font.color.rgb = RGBColor(0xA7, 0x31, 0x27)
    r.font.size = Pt(11)

    clausulas = _clausulas_iniciativa(iniciativa)
    doc.add_paragraph(clausulas["considerando"])

    doc.add_paragraph(
        "Considerando o término amigável da relação contratual, é "
        "celebrado o presente Termo de Distrato, regido pelas seguintes "
        "cláusulas e condições:"
    )

    # --- CLAUSULAS ---
    doc.add_paragraph()

    def _clausula(num: str, titulo: str, texto: str):
        p = doc.add_paragraph()
        r1 = p.add_run(f"CLÁUSULA {num} — {titulo}")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0xA7, 0x31, 0x27)
        doc.add_paragraph(texto)

    _clausula(
        "PRIMEIRA", "DO OBJETO",
        "O presente termo tem por objeto a RESCISÃO AMIGÁVEL do "
        "contrato de prestação de serviços contábeis anteriormente "
        "celebrado entre as partes, encerrando todas as obrigações "
        "decorrentes daquele instrumento, a partir da data de efeito "
        "abaixo registrada.",
    )

    _clausula(
        "SEGUNDA", "DA INICIATIVA",
        clausulas["clausula_motivacao"] +
        (f"\n\nMotivo registrado: {motivo}" if motivo else ""),
    )

    _clausula(
        "TERCEIRA", "DA DATA DE EFEITO",
        f"Este distrato produz efeitos a partir de "
        f"{_data_extenso(data_efeito)}, data em que cessam todas as "
        "obrigações da CONTRATADA com relação à escrita contábil, "
        "fiscal e trabalhista da CONTRATANTE.",
    )

    _clausula(
        "QUARTA", "DA DEVOLUÇÃO DOS DOCUMENTOS",
        "A CONTRATADA se compromete a devolver à CONTRATANTE, no prazo "
        "máximo de 15 (quinze) dias contados da assinatura deste termo, "
        "TODA a documentação contábil, fiscal, societária e trabalhista "
        "que esteja em sua posse, mediante recibo de entrega assinado "
        "pelo representante legal da CONTRATANTE.",
    )

    _clausula(
        "QUINTA", "DAS OBRIGAÇÕES PENDENTES",
        "As obrigações principais e acessórias com competência até a "
        "data de efeito deste distrato permanecem sob responsabilidade "
        "da CONTRATADA, que se compromete a entregar/transmitir todas "
        "as declarações, guias e demonstrações já vencidas ou em curso "
        "de vencimento dentro do prazo legal. As competências "
        "posteriores à data de efeito ficam sob responsabilidade da "
        "CONTRATANTE.",
    )

    _clausula(
        "SEXTA", "DOS HONORÁRIOS",
        "A CONTRATANTE declara estar quite com todos os honorários "
        "devidos à CONTRATADA até a data de efeito deste distrato, "
        "nada mais tendo a reclamar a esse título. Eventuais "
        "pendências financeiras anteriores deverão ser quitadas em "
        "até 30 (trinta) dias, sob pena de cobrança pelos meios "
        "legais cabíveis.",
    )

    _clausula(
        "SÉTIMA", "DO SIGILO PROFISSIONAL",
        "A CONTRATADA permanece obrigada ao dever de sigilo "
        "profissional sobre todas as informações da CONTRATANTE a "
        "que teve acesso durante a vigência do contrato, conforme "
        "art. 27 do Código de Ética Profissional do Contador "
        "(Resolução CFC 803/96).",
    )

    _clausula(
        "OITAVA", "DA QUITAÇÃO RECÍPROCA",
        "Com a assinatura deste termo e o cumprimento das obrigações "
        "nele previstas, as partes outorgam UMA À OUTRA a mais ampla, "
        "geral, plena, irrevogável e irretratável QUITAÇÃO de todas "
        "as obrigações decorrentes do contrato ora rescindido, para "
        "nada mais reclamar uma da outra, a qualquer título.",
    )

    _clausula(
        "NONA", "DO FORO",
        "Fica eleito o Foro da Comarca de Cotia/SP, com renúncia "
        "expressa a qualquer outro, por mais privilegiado que seja, "
        "para dirimir quaisquer questões decorrentes deste termo.",
    )

    # --- DATA E ASSINATURAS ---
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Cotia/SP, {_data_extenso(data_efeito)}.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Linhas de assinatura — CONTRATADA
    p = doc.add_paragraph("_" * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CSM CONTABILIDADE EMPRESARIAL LTDA")
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONTRATADA — Sócio(a) Administrador(a)")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()
    doc.add_paragraph()

    # Linhas de assinatura — CONTRATANTE
    p = doc.add_paragraph("_" * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(razao_cliente)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"CONTRATANTE — CNPJ {cnpj_cliente}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()
    doc.add_paragraph()

    # Testemunhas
    p = doc.add_paragraph()
    r = p.add_run("TESTEMUNHAS:")
    r.bold = True
    r.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph("1) " + "_" * 50 + "  CPF: " + "_" * 15)
    doc.add_paragraph()
    p = doc.add_paragraph("2) " + "_" * 50 + "  CPF: " + "_" * 15)

    # ============ Salva o Word ============
    doc.save(docx_path)

    # ============ Converte pra PDF (se solicitado) ============
    pdf_resultado = None
    if gerar_pdf:
        try:
            subprocess.run([
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", pasta_destino, docx_path,
            ], check=True, capture_output=True, timeout=60)
            if os.path.exists(pdf_path):
                pdf_resultado = pdf_path
        except Exception as exc:
            # PDF falhou mas o Word foi gerado — não bloqueia
            print(f"Falha ao gerar PDF: {exc}")

    return {
        "docx": docx_path,
        "pdf": pdf_resultado,
        "filename_base": base,
    }
